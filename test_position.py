from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz

doc = Document()
p = doc.add_paragraph()
p.add_run("Text before image ")

svg = '<svg width="100" height="100"><rect width="100" height="100" fill="red"/></svg>'
img_doc = fitz.open("svg", svg.encode("utf-8"))
pix = img_doc[0].get_pixmap()
run = p.add_run()
run.add_picture(pix.tobytes("png"), width=Pt(20))

rPr = run._element.get_or_add_rPr()
position = OxmlElement('w:position')
position.set(qn('w:val'), '-40') # Lower by 20 points
rPr.append(position)

p.add_run(" Text after image")
doc.save("test_position.docx")
