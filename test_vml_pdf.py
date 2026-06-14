from docx import Document
from docx.oxml import parse_xml
import os
import win32com.client

doc = Document()
doc.add_paragraph("This is a test document with VML watermark.")

for section in doc.sections:
    header = section.header
    p_watermark = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    run_watermark = p_watermark.add_run()
    xml = '''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">
        <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" path="m@7,l@8,m@5,21600l@6,21600e">
            <v:path textpathok="t"/>
            <v:textpath on="t" fitshape="t"/>
        </v:shapetype>
        <v:shape id="WordWaterMark" type="#_x0000_t136" style="position:absolute;left:0;margin-left:0;margin-top:0;width:500pt;height:200pt;rotation:315;z-index:-251658240;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin" fillcolor="#E0E0E0" stroked="f">
            <v:textpath on="t" style="font-family:'SimHei';font-size:60pt" string="融智云考题库"/>
        </v:shape>
    </w:pict>'''
    run_watermark._element.append(parse_xml(xml))

doc.save("test_vml_watermark.docx")

pdf_path = os.path.abspath("test_vml_watermark.pdf")
app = win32com.client.DispatchEx("Word.Application")
app.Visible = False
app.DisplayAlerts = 0

doc = app.Documents.Open(os.path.abspath("test_vml_watermark.docx"))
try:
    doc.SaveAs(pdf_path, 17)
    print("VML Watermark PDF Export Success!")
except Exception as e:
    print("VML Watermark PDF Export Failed:", e)
finally:
    doc.Close(0)
    app.Quit()
