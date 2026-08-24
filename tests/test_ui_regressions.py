import json
import os
import tempfile
import unittest
from unittest.mock import Mock, call, patch

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
os.environ.setdefault("QT_OPENGL", "software")
os.environ.setdefault(
    "QTWEBENGINE_CHROMIUM_FLAGS",
    "--disable-gpu --disable-gpu-compositing --disable-logging",
)

from PySide6.QtCore import QEvent, QPoint, QPointF, QUrl, Qt
from PySide6.QtGui import QCloseEvent, QMouseEvent
from PySide6.QtWebEngineCore import QWebEnginePage
from PySide6.QtWebEngineWidgets import QWebEngineView
from PySide6.QtWidgets import QApplication, QLineEdit, QMainWindow

from config import settings as config_settings
from config.settings import HARDCODED_SCHOOL_CODE
from config.version import APP_RELEASE, APP_VERSION, APP_VERSION_TUPLE
from modules.ai_answer import infer_answer_with_ai
from modules.exporter import (
    export_to_docx,
    export_to_markdown,
    export_to_txt,
    get_practice_line_count,
)
from ui.main_window import YunKaoExtractorApp
from ui.settings_dialog import SettingsDialog, extract_model_ids
from ui.title_bar import TitleBar


class UiRegressionTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.app = QApplication.instance() or QApplication([])

    def create_window(self, current_user="20260001"):
        load_patcher = patch.object(QWebEngineView, "load")
        load_patcher.start()
        self.addCleanup(load_patcher.stop)
        window = YunKaoExtractorApp(
            current_user=current_user,
            jwt_token="",
            user_data={"nickname": "测试用户"},
        )
        self.addCleanup(window.close)
        return window

    def test_config_migrates_legacy_user_to_shared_location(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            legacy_file = os.path.join(temp_dir, "legacy-config.json")
            shared_file = os.path.join(temp_dir, "shared", "config.json")
            with open(legacy_file, "w", encoding="utf-8") as handle:
                json.dump(
                    {
                        "config_version": 1,
                        "yunkao_user": "",
                        "user": "20260009",
                        "ai_provider": "custom",
                    },
                    handle,
                )

            with (
                patch.object(config_settings, "CONFIG_FILE", shared_file),
                patch.object(
                    config_settings,
                    "LEGACY_CONFIG_FILES",
                    (legacy_file,),
                    create=True,
                ),
            ):
                migrated = config_settings.load_config()

            self.assertEqual(migrated["yunkao_user"], "20260009")
            self.assertTrue(os.path.exists(shared_file))

    def test_main_window_restores_overlay_and_webchannel(self):
        window = self.create_window()

        self.assertTrue(hasattr(window, "overlay"))
        self.assertTrue(hasattr(window, "channel"))
        self.assertTrue(hasattr(window, "bridge"))
        self.assertFalse(window.overlay.isHidden())
        self.assertFalse(window.overlay.btn_settings.isHidden())
        self.assertFalse(window.overlay.btn_export.isHidden())
        self.assertEqual(window.overlay.btn_toggle.objectName(), "btn_primary")
        self.assertEqual(window.overlay.btn_export.objectName(), "btn_export")
        self.assertIn(APP_RELEASE, window.windowTitle())
        self.assertIn(APP_RELEASE, window.overlay.lbl_title.text())

        window.extracted_questions.append({"title": "测试题"})
        window.refresh_export_button()
        self.assertTrue(window.overlay.btn_export.isEnabled())

        window._set_extraction_ui(True)
        self.assertFalse(window.overlay.btn_export.isEnabled())
        self.assertTrue(window.overlay.btn_toggle.property("extracting"))
        self.assertEqual(window.overlay.btn_toggle.text(), "停止提取")

    def test_main_window_uses_docked_shell_with_legacy_overlay_alias(self):
        window = self.create_window()

        self.assertIs(window.overlay, window.control_panel)
        self.assertTrue(hasattr(window, "title_bar"))
        self.assertTrue(hasattr(window, "navigation"))
        self.assertTrue(hasattr(window, "browser_shell"))
        self.assertIs(window.browser, window.browser_shell.browser)
        self.assertTrue(window.navigation.buttons["workspace"].isChecked())
        self.assertEqual(window.overlay.btn_back, window.browser_shell.btn_back)

    def test_browser_fit_scales_wide_page_to_the_available_viewport(self):
        window = self.create_window()

        with patch.object(window.browser, "setZoomFactor") as set_zoom:
            window._apply_browser_fit({"viewport": 1000, "content": 1250})

        set_zoom.assert_called_once_with(0.8)

    def test_title_bar_drag_is_manual_and_does_not_trigger_windows_snap(self):
        window = QMainWindow()
        window.setWindowFlags(Qt.FramelessWindowHint | Qt.Window)
        window.resize(640, 480)
        window.move(100, 100)
        title_bar = TitleBar(window)
        title_bar.show()
        self.addCleanup(window.close)

        native_handle = Mock()
        native_handle.startSystemMove.return_value = True
        press_event = QMouseEvent(
            QEvent.MouseButtonPress,
            QPointF(20, 20),
            QPointF(20, 20),
            QPointF(120, 120),
            Qt.LeftButton,
            Qt.LeftButton,
            Qt.NoModifier,
        )
        move_event = QMouseEvent(
            QEvent.MouseMove,
            QPointF(60, 45),
            QPointF(60, 45),
            QPointF(160, 125),
            Qt.NoButton,
            Qt.LeftButton,
            Qt.NoModifier,
        )

        with patch.object(window, "windowHandle", return_value=native_handle):
            title_bar.mousePressEvent(press_event)
            title_bar.mouseMoveEvent(move_event)

        native_handle.startSystemMove.assert_not_called()
        self.assertEqual(window.pos(), QPoint(140, 105))

    def test_control_panel_metrics_and_export_state_are_synchronized(self):
        window = self.create_window()
        panel = window.control_panel

        panel.set_progress_text("进度: 12 / 80 (已存 11 题)")
        panel.set_run_metrics(ai_pending=2, average=1.6)
        self.assertEqual(panel.lbl_current.text(), "12")
        self.assertEqual(panel.lbl_total.text(), "/ 80 题")
        self.assertEqual(panel.lbl_saved.text(), "11")
        self.assertEqual(panel.lbl_ai.text(), "2")
        self.assertEqual(panel.lbl_average.text(), "1.6s")
        self.assertEqual(panel.progress_bar.value(), 15)

        panel.refresh_export_state(False)
        self.assertFalse(panel.btn_export.isEnabled())
        self.assertFalse(panel.btn_export_pdf.isEnabled())
        panel.refresh_export_state(True)
        self.assertTrue(panel.btn_export.isEnabled())
        self.assertTrue(panel.btn_export_more.isEnabled())

    def test_export_shortcuts_forward_their_explicit_format(self):
        window = self.create_window()
        export = Mock()
        window.export_basic_questions = export

        window.control_panel.btn_export_pdf.click()
        window.control_panel.btn_export.click()
        window.control_panel.btn_export_more.click()

        self.assertEqual(
            export.call_args_list,
            [call("PDF"), call("DOCX"), call(None)],
        )

    def test_direct_pdf_export_locks_dialog_filter_and_extension(self):
        window = self.create_window()
        window.extracted_questions.append({"title": "测试题"})

        export_thread = Mock()
        progress_dialog = Mock()
        with tempfile.TemporaryDirectory() as temp_dir:
            window.config["default_export_dir"] = temp_dir
            with (
                patch(
                    "ui.main_window.QFileDialog.getSaveFileName",
                    return_value=(os.path.join(temp_dir, "题库"), "PDF 文件 (*.pdf)"),
                ) as get_save_file_name,
                patch("ui.main_window.QProgressDialog", return_value=progress_dialog),
                patch("ui.main_window.ExportThread", return_value=export_thread) as export_cls,
            ):
                window.export_basic_questions("PDF")

        self.assertEqual(
            get_save_file_name.call_args.args[3],
            "PDF 文件 (*.pdf)",
        )
        self.assertTrue(export_cls.call_args.args[1].endswith(".pdf"))
        self.assertEqual(export_cls.call_args.args[2], "PDF 文件 (*.pdf)")

    def test_external_browser_button_opens_current_url(self):
        window = self.create_window()
        with (
            patch.object(
                window.browser,
                "url",
                return_value=QUrl("https://www.cctrcloud.net/practice/login.html"),
            ),
            patch("ui.main_window.QDesktopServices.openUrl", return_value=True) as open_url,
        ):
            window.browser_shell.btn_external.click()

        open_url.assert_called_once_with(
            QUrl("https://www.cctrcloud.net/practice/login.html")
        )

    def test_clearing_questions_keeps_running_ai_worker_referenced_until_finished(self):
        window = self.create_window()
        worker = Mock()
        worker.isRunning.return_value = True
        key = (window.ai_session_id, 0)
        window.pending_ai_workers[key] = worker
        window._live_ai_workers[key] = worker
        window.extracted_questions.append({"title": "待补全题"})

        window.clear_extracted_questions()

        self.assertNotIn(key, window.pending_ai_workers)
        self.assertIn(key, window._live_ai_workers)
        window._cleanup_ai_fill_worker(key)
        self.assertNotIn(key, window._live_ai_workers)
        worker.deleteLater.assert_called_once_with()

    def test_close_event_requests_running_ai_workers_before_destroying_window(self):
        window = self.create_window()
        worker = Mock()
        worker.isRunning.return_value = True
        window._live_ai_workers[(window.ai_session_id, 0)] = worker

        event = QCloseEvent()
        with patch("ui.main_window.QTimer.singleShot") as schedule:
            window.closeEvent(event)

        self.assertFalse(event.isAccepted())
        worker.requestInterruption.assert_called_once_with()
        self.assertTrue(window._close_after_ai_workers)
        schedule.assert_called_once()

        worker.isRunning.return_value = False
        window._close_after_ai_workers = False
        window.close()

    def test_save_config_never_persists_ai_key(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            config_file = os.path.join(temp_dir, "config.json")
            with patch.object(config_settings, "CONFIG_FILE", config_file):
                config_settings.save_config(
                    {"ai_api_key": "secret-key", "ai_key_saved": True, "value": 1}
                )
            with open(config_file, "r", encoding="utf-8") as handle:
                saved = json.load(handle)

        self.assertNotIn("ai_api_key", saved)
        self.assertTrue(saved["ai_key_saved"])

    def test_settings_open_does_not_fetch_models_without_explicit_click(self):
        config = {
            "default_export_dir": tempfile.gettempdir(),
            "default_filename_prefix": "题库",
            "yunkao_user": "",
            "ai_provider": "custom",
            "ai_base_url": "https://api.example.com/v1",
            "ai_model": "existing-model",
            "ai_api_key": "test-key",
        }
        with (
            patch("ui.settings_dialog.load_config", return_value=config),
            patch("ui.settings_dialog.keyring.get_password", return_value=None),
        ):
            dialog = SettingsDialog()
            self.addCleanup(dialog.close)
            with patch.object(dialog.network_manager, "get") as get_models:
                self.app.processEvents()
                self.assertFalse(get_models.called)

    def test_version_marker_is_semantic_and_visible_in_settings(self):
        self.assertRegex(APP_VERSION, r"^\d+\.\d+\.\d+$")
        self.assertEqual(len(APP_VERSION_TUPLE), 4)
        self.assertEqual(APP_VERSION_TUPLE[-1], 0)

        with (
            patch("ui.settings_dialog.load_config", return_value={}),
            patch("ui.settings_dialog.keyring.get_password", return_value=None),
        ):
            dialog = SettingsDialog()
            self.addCleanup(dialog.close)

        self.assertIn(APP_RELEASE, dialog.windowTitle())
        self.assertIn(APP_RELEASE, dialog.lbl_brand_subtitle.text())

    def test_overlay_expands_to_show_the_complete_action_row(self):
        window = self.create_window()
        overlay = window.overlay
        window.show()
        self.app.processEvents()
        overlay._refresh_expanded_size()
        self.app.processEvents()

        action_bottom = overlay.btn_export.mapTo(
            overlay,
            overlay.btn_export.rect().bottomRight(),
        ).y()
        self.assertGreaterEqual(
            overlay.height(),
            overlay.MIN_EXPANDED_HEIGHT,
        )
        self.assertLess(action_bottom, overlay.contentsRect().bottom())

        overlay.set_progress_text(
            "当前进度：正在处理包含图片选项和较长题目说明的练习内容"
        )
        overlay.set_mini_status(
            "题目内容正在准备，请稍候，完成后即可导出练习版文件"
        )
        self.app.processEvents()

        expanded_action_bottom = overlay.btn_export.mapTo(
            overlay,
            overlay.btn_export.rect().bottomRight(),
        ).y()
        self.assertLess(
            expanded_action_bottom,
            overlay.contentsRect().bottom(),
        )

    def test_settings_save_local_password_and_custom_api(self):
        with tempfile.TemporaryDirectory() as export_dir:
            config = {
                "default_export_dir": export_dir,
                "default_filename_prefix": "题库",
                "yunkao_user": "",
                "yunkao_remember_password": True,
                "ai_provider": "openai",
                "ai_base_url": "https://api.openai.com/v1",
                "ai_model": "gpt-4o-mini",
                "ai_api_key": "",
            }
            with (
                patch("ui.settings_dialog.load_config", return_value=config.copy()),
                patch("ui.settings_dialog.save_config") as save_config,
                patch("ui.settings_dialog.keyring.get_password", return_value=None),
                patch("ui.settings_dialog.keyring.set_password") as set_password,
            ):
                dialog = SettingsDialog()
                self.addCleanup(dialog.close)

                self.assertEqual(dialog.btn_toggle_yunkao_pwd.text(), "显示")
                dialog.btn_toggle_yunkao_pwd.setChecked(True)
                self.assertEqual(
                    dialog.txt_yunkao_pwd.echoMode(),
                    QLineEdit.Normal,
                )
                self.assertEqual(dialog.btn_toggle_yunkao_pwd.text(), "隐藏")
                dialog.txt_yunkao_user.setText("20260002")
                dialog.txt_yunkao_pwd.setText("local-secret")
                dialog.chk_remember_pwd.setChecked(True)
                dialog.txt_ai_url.setText("https://api.example.com/v1")
                dialog.txt_ai_model.setText("example-model")
                dialog.txt_ai_key.setText("example-key")
                dialog.save_settings()

            self.assertIn(
                call(
                    "YunKaoDesktop",
                    f"{HARDCODED_SCHOOL_CODE}_20260002",
                    "local-secret",
                ),
                set_password.call_args_list,
            )
            self.assertIn(
                call("YunKaoDesktop/AI", "openai", "example-key"),
                set_password.call_args_list,
            )
            saved = save_config.call_args.args[0]
            self.assertEqual(saved["yunkao_user"], "20260002")
            self.assertEqual(saved["ai_base_url"], "https://api.example.com/v1")
            self.assertEqual(saved["ai_model"], "example-model")
            self.assertNotIn("ai_api_key", saved)
            self.assertTrue(saved["ai_key_saved"])

    def test_custom_api_fetches_selectable_model_list(self):
        self.assertEqual(
            extract_model_ids({"data": [{"id": "model-a"}, {"id": "model-b"}]}),
            ["model-a", "model-b"],
        )
        self.assertEqual(
            extract_model_ids({"models": [{"name": "model-c"}, "model-d"]}),
            ["model-c", "model-d"],
        )

        with tempfile.TemporaryDirectory() as export_dir:
            config = {
                "default_export_dir": export_dir,
                "default_filename_prefix": "题库",
                "yunkao_user": "",
                "ai_provider": "custom",
                "ai_base_url": "https://api.example.com/v1/",
                "ai_model": "existing-model",
                "ai_api_key": "test-key",
            }
            with (
                patch("ui.settings_dialog.load_config", return_value=config),
                patch("ui.settings_dialog.keyring.get_password", return_value=None),
            ):
                dialog = SettingsDialog()
                self.addCleanup(dialog.close)

            self.assertEqual(dialog.btn_fetch_models.text(), "查询模型")
            self.assertIn("AI 生成内容可能不准确", dialog.lbl_ai_warning.text())
            reply = Mock()
            reply.finished = Mock()
            reply.finished.connect = Mock()
            with patch.object(dialog.network_manager, "get", return_value=reply) as get:
                dialog.fetch_models()

            request = get.call_args.args[0]
            self.assertEqual(
                request.url().toString(),
                "https://api.example.com/v1/models",
            )
            self.assertEqual(bytes(request.rawHeader("Authorization")), b"Bearer test-key")

            dialog._populate_models(["model-a", "model-b"])
            self.assertEqual(dialog.cmb_ai_model.count(), 2)
            self.assertEqual(dialog.cmb_ai_model.itemText(0), "model-a")

    def test_docx_marks_only_ai_generated_content(self):
        from docx import Document

        questions = [
            {
                "title": "普通题",
                "options": ["A. 正确"],
                "answer": "A",
                "answer_source": "dom",
                "analysis": "网页解析",
                "analysis_source": "dom",
            },
            {
                "title": "待补全题",
                "options": ["A. 正确", "B. 错误"],
                "answer": "B",
                "answer_source": "ai",
                "answer_confidence": 0.86,
                "analysis": "模型解析\n（AI生成）",
                "analysis_source": "ai",
            },
        ]

        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, "ai-marker.docx")
            export_to_docx(questions, file_path)
            paragraphs = [paragraph.text for paragraph in Document(file_path).paragraphs]

        ordinary_answer = next(text for text in paragraphs if text.startswith("[答案]:"))
        ai_answer = next(text for text in paragraphs if text.startswith("[答案 · AI生成"))
        ai_analysis = next(text for text in paragraphs if text.startswith("[解析 · AI生成]"))
        self.assertNotIn("AI生成", ordinary_answer)
        self.assertIn("置信度 0.86", ai_answer)
        self.assertEqual(ai_analysis.count("AI生成"), 1)

    def test_practice_export_hides_answers_and_uses_compact_spacing(self):
        from docx import Document

        questions = [
            {
                "title": "下列说法正确的是？",
                "question_type": "单选题",
                "options": ["A. 甲", "B. 乙"],
                "answer": "A",
                "analysis": "因为甲正确。",
            },
            {
                "title": "请简述测试的作用。",
                "question_type": "简答题",
                "options": [],
                "answer": "验证软件质量。",
                "analysis": "围绕验证展开。",
            },
        ]

        self.assertEqual(get_practice_line_count(questions[0]), 1)
        self.assertEqual(get_practice_line_count(questions[1]), 3)

        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = os.path.join(temp_dir, "practice.docx")
            markdown_path = os.path.join(temp_dir, "practice.md")
            txt_path = os.path.join(temp_dir, "practice.txt")

            export_to_docx(questions, docx_path, include_answers=False)
            export_to_markdown(questions, markdown_path, include_answers=False)
            export_to_txt(questions, txt_path, include_answers=False)

            document = Document(docx_path)
            docx_paragraphs = [
                paragraph.text for paragraph in document.paragraphs
            ]
            docx_text = "\n".join(docx_paragraphs)
            with open(markdown_path, "r", encoding="utf-8") as handle:
                markdown_text = handle.read()
            with open(txt_path, "r", encoding="utf-8") as handle:
                txt_text = handle.read()

        for exported_text in (docx_text, markdown_text, txt_text):
            self.assertNotIn("因为甲正确", exported_text)
            self.assertNotIn("验证软件质量", exported_text)
            self.assertNotIn("围绕验证展开", exported_text)
            self.assertIn("作答区", exported_text)

        self.assertIn("题库练习版", markdown_text)
        self.assertIn("题库练习版", txt_text)
        self.assertEqual(markdown_text.count("> ___"), 4)
        self.assertEqual(txt_text.count("  ___"), 4)
        self.assertNotIn("-" * 40, docx_paragraphs)
        self.assertNotIn("\n---\n", markdown_text)
        self.assertNotIn("-" * 30, txt_text)

    def test_docx_anti_resale_notice_only_appears_once(self):
        from docx import Document

        questions = [
            {
                "title": f"测试题 {index}",
                "question_type": "单选题",
                "options": ["A. 正确", "B. 错误"],
                "answer": "A",
            }
            for index in range(12)
        ]

        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, "notice-frequency.docx")
            export_to_docx(questions, file_path, watermark=False)
            document = Document(file_path)
            document_text = "\n".join(
                paragraph.text for paragraph in document.paragraphs
            )

        self.assertEqual(
            document_text.count("本资料由云考助手免费生成，禁止倒卖。"),
            1,
        )

    def test_practice_mode_toggle_updates_persisted_config(self):
        window = self.create_window()
        window.config["export_without_answers"] = False
        window.overlay.chk_practice_export.blockSignals(True)
        window.overlay.chk_practice_export.setChecked(False)
        window.overlay.chk_practice_export.blockSignals(False)

        with patch("ui.main_window.save_config") as save_config:
            window.overlay.chk_practice_export.setChecked(True)

        self.assertTrue(window.config["export_without_answers"])
        save_config.assert_called_once()
        self.assertIn("练习版", window.overlay.lbl_status_mini.text())

    def test_auto_fill_always_injects_school_and_user_without_password(self):
        window = self.create_window(current_user="20260003")
        scripts = []

        def capture_script(_page, script, *args):
            scripts.append(script)

        with (
            patch.object(
                QWebEngineView,
                "url",
                return_value=QUrl("https://www.cctrcloud.net/practice/login.html"),
            ),
            patch("ui.main_window.keyring.get_password", return_value=None),
            patch.object(QWebEnginePage, "runJavaScript", new=capture_script),
        ):
            window.trigger_auto_fill()

        self.assertEqual(len(scripts), 1)
        self.assertIn(HARDCODED_SCHOOL_CODE, scripts[0])
        self.assertIn("20260003", scripts[0])

    def test_auto_fill_with_empty_account_never_injects_local_user(self):
        window = self.create_window(current_user="")
        scripts = []

        def capture_script(_page, script, *args):
            scripts.append(script)

        with (
            patch.object(
                QWebEngineView,
                "url",
                return_value=QUrl("https://www.cctrcloud.net/practice/login.html"),
            ),
            patch("ui.main_window.keyring.get_password", return_value=None),
            patch.object(QWebEnginePage, "runJavaScript", new=capture_script),
        ):
            window.trigger_auto_fill()

        self.assertEqual(len(scripts), 1)
        self.assertIn(HARDCODED_SCHOOL_CODE, scripts[0])
        self.assertNotIn("local_user", scripts[0])

    def test_config_update_switches_active_student_number(self):
        window = self.create_window(current_user="old-user")
        with patch.object(window, "trigger_auto_fill") as trigger_auto_fill:
            window.update_config({"yunkao_user": "new-user"})

        self.assertEqual(window.current_user, "new-user")
        trigger_auto_fill.assert_called_once_with()

    def test_custom_api_configuration_drives_direct_request(self):
        response = Mock()
        response.raise_for_status.return_value = None
        response.json.return_value = {
            "choices": [
                {
                    "message": {
                        "content": '{"answer":"A","analysis":"解析","confidence":0.9}'
                    }
                }
            ],
            "usage": {"total_tokens": 12},
        }
        config = {
            "ai_provider": "custom",
            "ai_base_url": "https://api.example.com/v1",
            "ai_api_key": "test-key",
            "ai_model": "test-model",
            "ai_supports_images": False,
        }

        with patch("modules.ai_answer.requests.post", return_value=response) as post:
            result = infer_answer_with_ai(
                {"question_type": "单选题", "title": "测试", "options": ["A. 是"]},
                config,
            )

        request = post.call_args
        self.assertEqual(request.args[0], "https://api.example.com/v1/chat/completions")
        self.assertEqual(request.kwargs["headers"]["Authorization"], "Bearer test-key")
        self.assertEqual(request.kwargs["json"]["model"], "test-model")
        self.assertEqual(result["answer"], "A")


if __name__ == "__main__":
    unittest.main()
