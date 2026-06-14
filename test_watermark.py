from docx import Document
from docx.oxml import parse_xml

def add_watermark(doc, text="融智云考题库"):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = p.add_run()
        
        xml = f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:v="urn:schemas-microsoft-com:vml">
            <v:shape id="WordWaterMark" style="position:absolute;left:0;margin-left:0;margin-top:0;width:500pt;height:200pt;rotation:315;z-index:-251658240;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin" fillcolor="#D0D0D0" stroked="f">
                <v:fill opacity="0.2"/>
                <v:textpath style="font-family:'SimHei';font-size:60pt" string="{text}"/>
            </v:shape>
        </w:pict>'''
        pict = parse_xml(xml)
        run._element.append(pict)

doc = Document()
doc.add_paragraph("Test Document")
add_watermark(doc, "融智云考题库")
doc.save("test_watermark.docx")
