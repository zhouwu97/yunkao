import os
from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from PIL import Image, ImageDraw, ImageFont
import win32com.client

# 1. Generate watermark image
img = Image.new('RGBA', (800, 800), (255, 255, 255, 0))
d = ImageDraw.Draw(img)
# Since we might not know the exact font path, use default
try:
    font = ImageFont.truetype("msyh.ttc", 80)
except:
    font = ImageFont.load_default()
d.text((100, 400), "融智云考题库", fill=(200, 200, 200, 128), font=font)
img = img.rotate(45, expand=1)
img.save("watermark_temp.png")

# 2. Create document and insert image into header
doc = Document()
doc.add_paragraph("This is the document body. " * 50)

header = doc.sections[0].header
p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
run = p.add_run()
picture = run.add_picture("watermark_temp.png", width=Inches(6.0))

# Convert wp:inline to wp:anchor to float it behind text
drawing = picture._inline.getparent()
xml = drawing.xml
xml = xml.replace('wp:inline', 'wp:anchor')
xml = xml.replace('</wp:inline>', '</wp:anchor>')
# Add behindDoc="1"
xml = xml.replace('<wp:anchor', '<wp:anchor behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" relativeHeight="0"')
# Replace extents with positioning
import re
# Remove existing simplePos
xml = re.sub(r'<wp:simplePos[^>]*/>', '', xml)
# Add position elements
pos_xml = '''
<wp:simplePos x="0" y="0"/>
<wp:positionH relativeFrom="margin">
    <wp:align>center</wp:align>
</wp:positionH>
<wp:positionV relativeFrom="margin">
    <wp:align>center</wp:align>
</wp:positionV>
'''
xml = xml.replace('<wp:extent', pos_xml + '<wp:extent')

# Apply back
new_drawing = parse_xml(xml)
drawing.getparent().replace(drawing, new_drawing)

doc.save("test_drawingml_watermark.docx")

# 3. Test PDF export
pdf_path = os.path.abspath("test_drawingml_watermark.pdf")
app = win32com.client.DispatchEx("Word.Application")
app.Visible = False
app.DisplayAlerts = 0

word_doc = app.Documents.Open(os.path.abspath("test_drawingml_watermark.docx"))
try:
    word_doc.SaveAs(pdf_path, 17)
    print("DrawingML Watermark PDF Export Success!")
except Exception as e:
    print("DrawingML Watermark PDF Export Failed:", e)
finally:
    word_doc.Close(0)
    app.Quit()
