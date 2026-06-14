from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import math

doc = Document()
p = doc.add_paragraph()
p.add_run("Normal text here ")

run = p.add_run("LOWERED TEXT ")
rPr = run._element.get_or_add_rPr()
position = OxmlElement('w:position')
position.set(qn('w:val'), '-12')  # -6 points
rPr.append(position)

p.add_run(" and normal again.")
doc.save("e:/AI/yunkao/test_lower.docx")
