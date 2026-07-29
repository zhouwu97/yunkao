import re


MAX_IMAGE_WIDTH_PT = 420.0
MAX_IMAGE_HEIGHT_PT = 620.0
ANTI_RESALE_TEXT = "本资料由云考助手免费生成，禁止倒卖。\n如果你是付费购买所得，请立即退款。\n官方免费获取：GitHub zhouwu97 / SYLUlive"
PDF_PROTECTION_MODE = "normal"  # normal / flatten
PDF_FLATTEN_DPI = 130

SCORE_MARKER_PATTERN = re.compile(r"[ \t\u00a0]{2,}([（(]?\d+\s*分[）)]?)")
LEGACY_AI_SUFFIX_PATTERN = re.compile(r"\s*[（(]AI\s*生成[）)]\s*$", re.IGNORECASE)



def normalize_export_text_run(text):
    """Normalize web layout whitespace before writing text into Word runs."""
    text = text.replace("\u00a0", " ")
    return SCORE_MARKER_PATTERN.sub(r" \1", text)


def should_add_space_before_inline_image(text):
    """Add a separator when text such as '0.5' is immediately followed by math."""
    text = text.replace("\u00a0", " ")
    if not text or text[-1].isspace():
        return False
    return bool(re.search(r"[A-Za-z0-9)\]）}<>≤≥=]$", text))


def calculate_contained_point_size(width_pt, height_pt):
    """Constrain an already point-sized object to the export page bounds."""
    if width_pt <= 0 or height_pt <= 0:
        return None, None
    scale = min(
        1.0,
        MAX_IMAGE_WIDTH_PT / width_pt,
        MAX_IMAGE_HEIGHT_PT / height_pt,
    )
    return width_pt * scale, height_pt * scale


def calculate_contained_image_size(width_px, height_px, requested_width_pt=None,
                                   requested_height_pt=None):
    """Returns a page-safe image size while preserving the source aspect ratio."""
    if width_px <= 0 or height_px <= 0:
        return None, None

    source_width_pt = width_px * 0.75
    source_height_pt = height_px * 0.75
    width_pt = requested_width_pt or source_width_pt
    height_pt = requested_height_pt or source_height_pt

    if requested_width_pt is not None and requested_height_pt is None:
        height_pt = requested_width_pt * height_px / width_px
    elif requested_height_pt is not None and requested_width_pt is None:
        width_pt = requested_height_pt * width_px / height_px

    scale = min(
        1.0,
        MAX_IMAGE_WIDTH_PT / width_pt,
        MAX_IMAGE_HEIGHT_PT / height_pt,
    )
    return width_pt * scale, height_pt * scale


def get_export_label(question, content_type):
    """根据内容来源生成统一标签，确保各导出格式能识别 AI 内容。"""
    if content_type == "answer":
        label = "答案"
        source = question.get("answer_source")
    else:
        label = "解析"
        source = question.get("analysis_source")

    if source != "ai":
        return f"[{label}]"

    confidence = question.get("answer_confidence") if content_type == "answer" else None
    if isinstance(confidence, (int, float)):
        return f"[{label} · AI生成 | 置信度 {confidence:.2f}]"
    return f"[{label} · AI生成]"


def normalize_ai_content(text, source):
    """移除旧数据中的尾注，避免与结构化 AI 标签重复。"""
    if source == "ai":
        return LEGACY_AI_SUFFIX_PATTERN.sub("", text)
    return text


def get_practice_line_count(question):
    """按题型返回紧凑作答行数，避免练习版出现过多空白。"""
    question_type = str(question.get("question_type", "") or "")
    compact_keywords = ("选择", "判断", "填空")
    if any(keyword in question_type for keyword in compact_keywords):
        return 1
    if question.get("options"):
        return 1
    return 3


def _write_markdown_practice_area(file_handle, question):
    file_handle.write("\n**作答区**\n")
    for _ in range(get_practice_line_count(question)):
        file_handle.write("> ________________________________________________\n")


def _write_txt_practice_area(file_handle, question):
    file_handle.write("\n  [作答区]\n")
    for _ in range(get_practice_line_count(question)):
        file_handle.write("  ________________________________________________\n")


def export_to_markdown(questions, file_path, include_answers=True):
    with open(file_path, 'w', encoding='utf-8') as f:
        title = "题库导出" if include_answers else "题库练习版"
        f.write(f"# {title}\n\n")
        for i, q in enumerate(questions, 1):
            f.write(f"### {i}. {q['title']}\n\n")
            for opt in q['options']:
                f.write(f"- {opt}\n")
            
            if not include_answers:
                _write_markdown_practice_area(f, q)
            elif q.get('answer'):
                answer_title = get_export_label(q, "answer")
                f.write(f"\n**{answer_title}**\n{q['answer']}\n")
            if include_answers and q.get('analysis'):
                analysis_title = get_export_label(q, "analysis")
                analysis = normalize_ai_content(q['analysis'], q.get("analysis_source"))
                f.write(f"\n**{analysis_title}**\n{analysis}\n")
                
            # 练习版依靠作答线和自然留白分隔题目，避免打印出突兀的黑色横线。
            f.write("\n---\n\n" if include_answers else "\n\n")

def export_to_txt(questions, file_path, include_answers=True):
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write("题库导出\n" if include_answers else "题库练习版\n")
        f.write("="*30 + "\n\n")
        for i, q in enumerate(questions, 1):
            # TXT 去除图片 markdown 标记
            title = re.sub(r'!\[[^\]]*\]\((?:[^)(]+|\([^)(]*\))*\)', '[图片]', q['title'])
            f.write(f"{i}. {title}\n")
            for opt in q['options']:
                opt_txt = re.sub(r'!\[[^\]]*\]\((?:[^)(]+|\([^)(]*\))*\)', '[图片]', opt)
                f.write(f"  {opt_txt}\n")
                
            if not include_answers:
                _write_txt_practice_area(f, q)
            elif q.get('answer'):
                ans_txt = re.sub(r'!\[[^\]]*\]\((?:[^)(]+|\([^)(]*\))*\)', '[图片]', q['answer'])
                answer_title = get_export_label(q, "answer")
                f.write(f"\n  {answer_title}: {ans_txt}\n")
            if include_answers and q.get('analysis'):
                analysis = normalize_ai_content(q['analysis'], q.get("analysis_source"))
                ana_txt = re.sub(r'!\[[^\]]*\]\((?:[^)(]+|\([^)(]*\))*\)', '[图片]', analysis)
                f.write(f"  {get_export_label(q, 'analysis')}: {ana_txt}\n")
                
            if include_answers:
                f.write("\n" + "-"*30 + "\n\n")
            else:
                f.write("\n\n")

def export_to_docx(questions, file_path, progress_callback=None, watermark=True,
                   include_answers=True):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
    except ImportError:
        raise ImportError("未安装 python-docx 库，无法导出 Word 文档。请运行 pip install python-docx")

    import re
    import requests
    from io import BytesIO
    from docx.text.paragraph import Paragraph
    from docx.oxml import OxmlElement

    def insert_paragraph_after(paragraph):
        new_element = OxmlElement("w:p")
        paragraph._p.addnext(new_element)
        return Paragraph(new_element, paragraph._parent)

    def add_rich_text_to_paragraph(p, text):
        """将包含 ![img](url) 的富文本添加到段落中，自动下载并嵌入图片"""
        current_p = p
        last_end = 0
        for match in re.finditer(r'!\[[^\]]*\]<([^>]+)>', text):
            # 添加图片前面的文本
            text_before = text[last_end:match.start()]
            if text_before:
                normalized_before = normalize_export_text_run(text_before)
                current_p.add_run(normalized_before)
                if should_add_space_before_inline_image(normalized_before):
                    current_p.add_run(" ")
            
            # 尝试下载并嵌入图片
            raw_url = match.group(1)
            align_val, align_unit = None, None
            explicit_w_val, explicit_w_unit = None, None
            explicit_h_val, explicit_h_unit = None, None
            
            url = raw_url.split('|align:')[0].split('|w:')[0].split('|h:')[0]
            
            a_match = re.search(r'\|align:([-0-9.]+)(px|ex|em|pt)', raw_url)
            if a_match: align_val, align_unit = float(a_match.group(1)), a_match.group(2)
                
            w_match = re.search(r'\|w:([-0-9.]+)(px|ex|em|pt|%)', raw_url)
            if w_match: explicit_w_val, explicit_w_unit = float(w_match.group(1)), w_match.group(2)

            h_match = re.search(r'\|h:([-0-9.]+)(px|ex|em|pt|%)', raw_url)
            if h_match: explicit_h_val, explicit_h_unit = float(h_match.group(1)), h_match.group(2)

            def to_points(value, unit, maximum):
                if value is None:
                    return None
                if unit == 'px':
                    return value * 0.75
                if unit == 'pt':
                    return value
                if unit == 'ex':
                    return value * 8.0
                if unit == 'em':
                    return value * 16.0
                if unit == '%':
                    return maximum * value / 100.0
                return None

            explicit_width_pt = to_points(
                explicit_w_val, explicit_w_unit, MAX_IMAGE_WIDTH_PT
            )
            explicit_height_pt = to_points(
                explicit_h_val, explicit_h_unit, MAX_IMAGE_HEIGHT_PT
            )
                    
            def apply_w_position(run_obj, a_val, a_unit):
                if a_val is None: return
                offset_half_pt = 0
                if a_unit == 'px': offset_half_pt = a_val * 1.5
                elif a_unit == 'pt': offset_half_pt = a_val * 2
                elif a_unit == 'ex': offset_half_pt = a_val * 12
                elif a_unit == 'em': offset_half_pt = a_val * 24
                
                if offset_half_pt != 0:
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    rPr = run_obj._element.get_or_add_rPr()
                    pos = OxmlElement('w:position')
                    pos.set(qn('w:val'), str(int(offset_half_pt)))
                    rPr.append(pos)
            try:
                # 补全相对路径和协议头
                if url.startswith("//"):
                    url = "https:" + url
                elif url.startswith("/"):
                    url = "https://www.cctrcloud.net" + url
                elif not url.startswith("http") and not url.startswith("data:"):
                    # 可能是未知格式，暂时忽略
                    pass
                
                if url.startswith("data:image/svg+xml"):
                    import urllib.parse
                    import fitz  # PyMuPDF
                    
                    prefix, data = url.split(',', 1)
                    if ';base64' in prefix:
                        import base64
                        data += '=' * (-len(data) % 4)
                        svg_content = base64.b64decode(data).decode('utf-8')
                    else:
                        svg_content = urllib.parse.unquote(data)
                        
                    try:
                        # 使用 PyMuPDF 将 SVG 渲染为 PNG
                        # 强力修复 MathJax SVG 尺寸过大问题 (PyMuPDF 无法识别 ex/em 单位，会回退到 viewBox 导致巨大化)
                        svg_width_pt, svg_height_pt = None, None
                        
                        try:
                            w_match = re.search(r'width="([0-9.]+)(ex|em|px|pt)"', svg_content)
                            h_match = re.search(r'height="([0-9.]+)(ex|em|px|pt)"', svg_content)
                            if w_match:
                                val, unit = float(w_match.group(1)), w_match.group(2)
                                if unit == 'ex': svg_width_pt = val * 8.0
                                elif unit == 'em': svg_width_pt = val * 16.0
                                elif unit == 'px': svg_width_pt = val * 0.75
                                elif unit == 'pt': svg_width_pt = val
                            if h_match:
                                val, unit = float(h_match.group(1)), h_match.group(2)
                                if unit == 'ex': svg_height_pt = val * 8.0
                                elif unit == 'em': svg_height_pt = val * 16.0
                                elif unit == 'px': svg_height_pt = val * 0.75
                                elif unit == 'pt': svg_height_pt = val
                        except Exception as e:
                            pass
                            
                        # 获取 SVG 原始定义的物理宽高 (作为备选)
                        doc = fitz.open("svg", svg_content.encode('utf-8'))
                        if svg_width_pt is None: svg_width_pt = doc[0].rect.width
                        if svg_height_pt is None: svg_height_pt = doc[0].rect.height
                        
                        # 高清渲染 (300dpi)
                        pix = doc[0].get_pixmap(alpha=True, dpi=300)
                        image_stream = BytesIO(pix.tobytes("png"))
                        run = current_p.add_run()
                        image_width_pt, image_height_pt = calculate_contained_point_size(
                            svg_width_pt,
                            svg_height_pt,
                        )
                        run.add_picture(
                            image_stream,
                            width=Pt(image_width_pt),
                            height=Pt(image_height_pt),
                        )
                        
                        # 优先使用 HTML style 提取到的对齐信息
                        if align_val is not None:
                            apply_w_position(run, align_val, align_unit)
                        else:
                            # 尝试解析 SVG 的 vertical-align 以精准对齐 baseline
                            try:
                                v_align_match = re.search(r'vertical-align:\s*([-0-9.]+)ex', svg_content)
                                height_match = re.search(r'height="([0-9.]+)ex"', svg_content)
                                if v_align_match and height_match:
                                    v_align_ex = float(v_align_match.group(1))
                                    height_ex = float(height_match.group(1))
                                    if height_ex > 0:
                                        offset_pt = (v_align_ex / height_ex) * svg_height_pt
                                        from docx.oxml import OxmlElement
                                        from docx.oxml.ns import qn
                                        rPr = run._element.get_or_add_rPr()
                                        position = OxmlElement('w:position')
                                        position.set(qn('w:val'), str(int(offset_pt * 2)))
                                        rPr.append(position)
                            except Exception as e_offset:
                                print("Failed to apply vertical-align offset:", e_offset)
                            
                    except Exception as e:
                        print(f"PyMuPDF failed to render SVG: {e}")
                        current_p.add_run("[图片加载失败]")
                    
                    last_end = match.end()
                    continue

                # 处理普通的 base64 图片
                elif url.startswith("data:image"):
                    import base64
                    header, data = url.split(',', 1)
                    data += '=' * (-len(data) % 4)  # 修复 Incorrect padding
                    image_data = base64.b64decode(data)
                    image_stream = BytesIO(image_data)
                    from PIL import Image
                    img = Image.open(BytesIO(image_data))
                    width, height = img.size
                    block_image = align_val is None and (width > 150 or height > 100)
                    if block_image:
                        current_p = insert_paragraph_after(current_p)
                    run = current_p.add_run()
                    image_width, image_height = calculate_contained_image_size(
                        width, height, explicit_width_pt, explicit_height_pt
                    )
                    run.add_picture(
                        image_stream,
                        width=Pt(image_width),
                        height=Pt(image_height),
                    )
                    if align_val is not None: apply_w_position(run, align_val, align_unit)
                    if block_image:
                        current_p = insert_paragraph_after(current_p)
                    last_end = match.end()
                    continue

                response = requests.get(url, timeout=5)
                if response.status_code == 200:
                    from PIL import Image
                    image_data = response.content
                    image_stream = BytesIO(image_data)
                    
                    try:
                        img = Image.open(BytesIO(image_data))
                        width, height = img.size
                        
                        is_large = width > 150 or height > 100
                        if is_large and align_val is None:
                            current_p = insert_paragraph_after(current_p)
                                
                        run = current_p.add_run()
                        
                        image_width, image_height = calculate_contained_image_size(
                            width, height, explicit_width_pt, explicit_height_pt
                        )
                        run.add_picture(
                            image_stream,
                            width=Pt(image_width),
                            height=Pt(image_height),
                        )
                            
                        if align_val is not None: apply_w_position(run, align_val, align_unit)
                        if is_large and align_val is None:
                            current_p = insert_paragraph_after(current_p)
                    except Exception:
                        run = current_p.add_run()
                        run.add_picture(image_stream)
                        if align_val is not None: apply_w_position(run, align_val, align_unit)
                else:
                    current_p.add_run("[图片加载失败]")
            except Exception as e:
                print(f"Failed to load image: {url}, Error: {e}")
                current_p.add_run("[图片加载失败]")
                
            last_end = match.end()
            
        # 添加最后剩余的文本
        text_after = text[last_end:]
        if text_after:
            current_p.add_run(normalize_export_text_run(text_after))

    def add_practice_area(question):
        """使用浅色书写线生成克制的练习作答区。"""
        from docx.enum.text import WD_LINE_SPACING

        label = doc.add_paragraph()
        label.paragraph_format.space_before = Pt(3)
        label.paragraph_format.space_after = Pt(1)
        label_run = label.add_run("作答区")
        label_run.bold = True
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = RGBColor(80, 96, 112)

        line_count = get_practice_line_count(question)
        for line_index in range(line_count):
            line = doc.add_paragraph()
            line.paragraph_format.space_before = Pt(0)
            line.paragraph_format.space_after = Pt(
                8 if line_index == line_count - 1 else 2
            )
            line.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            line.paragraph_format.line_spacing = Pt(17)

            paragraph_properties = line._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "C8D2DE")
            borders.append(bottom)
            paragraph_properties.append(borders)

    doc = Document()
    
    # 强制全局使用宋体，避免 WPS 在包含图片的段落中因为渲染路径不同导致字体发粗（假黑体）的 Bug
    from docx.oxml.ns import qn
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置标题
    document_title = "融智云考题库导出" if include_answers else "融智云考题库练习版"
    heading = doc.add_heading(document_title, 0)
    heading.alignment = 1 # 居中
    
    # 开头添加防诈骗提示
    anti_scam_text = ANTI_RESALE_TEXT
    p_scam_start = doc.add_paragraph()
    p_scam_start.alignment = 1
    run_scam_start = p_scam_start.add_run(anti_scam_text)
    run_scam_start.font.color.rgb = RGBColor(255, 0, 0)
    
    total = len(questions)
    for i, q in enumerate(questions, 1):
        if progress_callback:
            progress_callback(i, total, f"正在处理第 {i}/{total} 题，渲染图片公式中...")
        # 题目
        p = doc.add_paragraph()
        p.add_run(f"{i}. ")
        add_rich_text_to_paragraph(p, q['title'])
        # 选项
        for opt in q['options']:
            p_opt = doc.add_paragraph(style='List Bullet')
            add_rich_text_to_paragraph(p_opt, opt)
            
        # 答案
        if not include_answers:
            add_practice_area(q)
        elif q.get('answer'):
            p_ans = doc.add_paragraph()
            run = p_ans.add_run(f"{get_export_label(q, 'answer')}: ")
            run.font.color.rgb = RGBColor(0, 112, 192) # 蓝色
            add_rich_text_to_paragraph(p_ans, q['answer'])
            
        # 解析
        if include_answers and q.get('analysis'):
            p_ana = doc.add_paragraph()
            run = p_ana.add_run(f"{get_export_label(q, 'analysis')}: ")
            run.font.color.rgb = RGBColor(237, 125, 49) # 橙色
            analysis = normalize_ai_content(q['analysis'], q.get("analysis_source"))
            add_rich_text_to_paragraph(p_ana, analysis)
            
        if include_answers:
            doc.add_paragraph("-" * 40)

    # 正文提示仅在文档开头出现一次；页内保护交给低透明度斜向水印，
    # 避免重复提示挤占题目和手写作答空间。
        
    # 添加全页水印
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    if watermark:
        for section in doc.sections:
            header = section.header
            p_watermark = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            run_watermark = p_watermark.add_run()
            xml = f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">
                <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" path="m@7,l@8,m@5,21600l@6,21600e">
                    <v:path textpathok="t"/>
                    <v:textpath on="t" fitshape="t"/>
                </v:shapetype>
                <v:shape id="WordWaterMark1" type="#_x0000_t136" style="position:absolute;left:0;margin-left:0;margin-top:-100pt;width:600pt;height:150pt;rotation:315;z-index:-251658240;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin" fillcolor="#E0E0E0" stroked="f">
                    <v:textpath on="t" style="font-family:'SimHei';font-size:50pt" string="免费题库，请勿购买"/>
                </v:shape>
                <v:shape id="WordWaterMark2" type="#_x0000_t136" style="position:absolute;left:0;margin-left:0;margin-top:100pt;width:600pt;height:100pt;rotation:315;z-index:-251658240;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin" fillcolor="#E0E0E0" stroked="f">
                    <v:textpath on="t" style="font-family:'Arial';font-size:40pt" string="GitHub zhouwu97 / SYLUlive"/>
                </v:shape>
            </w:pict>'''
            run_watermark._element.append(parse_xml(xml))
        
    doc.save(file_path)

def export_to_pdf(questions, file_path, progress_callback=None, include_answers=True):
    import os
    import shutil
    import tempfile

    temp_dir = tempfile.mkdtemp(prefix="yunkao_pdf_")
    temp_docx = os.path.join(temp_dir, "source.docx")
    temp_pdf = os.path.join(temp_dir, "converted.pdf")

    if os.name == "nt":
        try:
            import ctypes
            ctypes.windll.kernel32.SetFileAttributesW(temp_dir, 0x02)
        except Exception:
            pass
    
    try:
        export_to_docx(
            questions,
            temp_docx,
            progress_callback,
            watermark=False,
            include_answers=include_answers,
        )
        if progress_callback:
            progress_callback(len(questions), len(questions), "正在启动 Office 引擎生成 PDF (兼容 WPS)...")
            
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        
        try:
            # 优先尝试 WPS，防止用户电脑上有过期/损坏的微软 Office 导致命令拒绝
            word = None
            for app_name in ["KWPS.Application", "WPS.Application", "Word.Application"]:
                try:
                    word = win32com.client.DispatchEx(app_name)
                    break
                except:
                    continue
            
            if not word:
                raise Exception("未找到可用的 Word 或 WPS 组件")
                
            word.Visible = False
            word.DisplayAlerts = 0
            try:
                word.ScreenUpdating = False
            except Exception:
                pass

            try:
                doc = word.Documents.Open(
                    os.path.abspath(temp_docx),
                    False,
                    True,
                    False,
                )
            except Exception:
                doc = word.Documents.Open(os.path.abspath(temp_docx))
            try:
                doc.Windows(1).Visible = False
            except Exception:
                pass
            try:
                # 导出到临时 PDF，避免 WPS 云同步功能干扰或增量保存导致文件丢失
                if os.path.exists(temp_pdf):
                    os.remove(temp_pdf)
                        
                try:
                    doc.SaveAs(os.path.abspath(temp_pdf), 17) # 17 = wdFormatPDF
                except Exception as save_err:
                    try:
                        # 兼容部分极其特殊的 WPS 版本
                        doc.ExportAsFixedFormat(os.path.abspath(temp_pdf), 17)
                    except:
                        raise save_err
            finally:
                try:
                    doc.Close(0)
                finally:
                    word.Quit()
        except Exception as e_com:
            raise RuntimeError(f"调用 Office 组件失败，请确保没有打开同名 PDF 文件，且 WPS 处于正常状态: {str(e_com)}")
        finally:
            pythoncom.CoUninitialize()

        # 4. 成功生成 PDF 后，使用 PyMuPDF 添加防盗版水印 (绕过 WPS 渲染 VML 崩溃的 BUG)
        if progress_callback:
            progress_callback(len(questions), len(questions), "正在添加防盗版水印...")
        import fitz
        from PIL import Image, ImageDraw, ImageFont
        import io
        
        img = Image.new('RGBA', (800, 800), (255, 255, 255, 0))
        d = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("C:\\Windows\\Fonts\\msyh.ttc", 50)
            font_small = ImageFont.truetype("C:\\Windows\\Fonts\\msyh.ttc", 40)
        except:
            font = ImageFont.load_default()
            font_small = ImageFont.load_default()
            
        text1 = "免费题库，请勿购买"
        text2 = "官方免费获取：GitHub zhouwu97 / SYLUlive"
        
        def draw_centered_text(draw, txt, y_offset, fnt):
            try:
                bbox = draw.textbbox((0, 0), txt, font=fnt)
                w = bbox[2] - bbox[0]
                h = bbox[3] - bbox[1]
            except AttributeError:
                w, h = draw.textsize(txt, font=fnt)
            x = (800 - w) / 2
            y = (800 - h) / 2 + y_offset
            draw.text((x, y), txt, fill=(200, 200, 200, 80), font=fnt)
            
        draw_centered_text(d, text1, -40, font)
        draw_centered_text(d, text2, 40, font_small)
        img = img.rotate(45, resample=Image.BICUBIC)
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_bytes = img_byte_arr.getvalue()
        
        # 确保目标文件未被占用
        target_pdf = os.path.abspath(file_path)
        if os.path.exists(target_pdf):
            try:
                os.remove(target_pdf)
            except Exception as rm_err:
                raise RuntimeError(f"无法覆盖已存在的文件，请确保该 PDF 未在其他软件中打开！({str(rm_err)})")

        import secrets
        owner_pw = secrets.token_hex(20)
        # 允许打印、复制、高质量打印和辅助功能。禁止：修改、批注、组装、表单填写
        perm = int(fitz.PDF_PERM_PRINT | fitz.PDF_PERM_COPY | fitz.PDF_PERM_ACCESSIBILITY | fitz.PDF_PERM_PRINT_HQ)
        def _save_text_pdf_with_watermark(pdf_doc, target_pdf, watermark_bytes, owner_pw, permissions):
            import fitz
            for page in pdf_doc:
                rect = page.rect
                watermark_rect = fitz.Rect(
                    (rect.width - 360) / 2,
                    (rect.height - 360) / 2 + 20,
                    (rect.width + 360) / 2,
                    (rect.height + 360) / 2 + 20,
                )
                page.insert_image(watermark_rect, stream=watermark_bytes, overlay=True)

            pdf_doc.save(
                target_pdf,
                encryption=fitz.PDF_ENCRYPT_AES_256,
                owner_pw=owner_pw,
                permissions=permissions,
                garbage=4,
                deflate=True,
                clean=True,
            )

        def _save_flattened_pdf_with_watermark(pdf_doc, target_pdf, watermark_bytes, owner_pw, permissions):
            import fitz
            for page in pdf_doc:
                rect = page.rect
                watermark_rect = fitz.Rect(
                    (rect.width - 360) / 2,
                    (rect.height - 360) / 2 + 20,
                    (rect.width + 360) / 2,
                    (rect.height + 360) / 2 + 20,
                )
                page.insert_image(watermark_rect, stream=watermark_bytes, overlay=True)

            img_pdf = fitz.open()
            for page in pdf_doc:
                pix = page.get_pixmap(dpi=PDF_FLATTEN_DPI, alpha=False)
                new_page = img_pdf.new_page(width=page.rect.width, height=page.rect.height)
                new_page.insert_image(new_page.rect, stream=pix.tobytes("png"))
                
            img_pdf.save(
                target_pdf,
                encryption=fitz.PDF_ENCRYPT_AES_256,
                owner_pw=owner_pw,
                permissions=permissions,
                garbage=4,
                deflate=True,
                clean=True,
            )
            img_pdf.close()

        pdf_doc = fitz.open(temp_pdf)
        try:
            if PDF_PROTECTION_MODE == "flatten":
                _save_flattened_pdf_with_watermark(pdf_doc, target_pdf, img_bytes, owner_pw, perm)
            else:
                _save_text_pdf_with_watermark(pdf_doc, target_pdf, img_bytes, owner_pw, perm)
        finally:
            pdf_doc.close()
            
    except Exception as e:
        raise RuntimeError(f"PDF 转换失败: {str(e)}")
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)
